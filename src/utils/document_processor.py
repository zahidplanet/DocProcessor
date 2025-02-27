"""
Utilities for processing different document types.
"""

import os
import PyPDF2
import docx
import openpyxl
import pandas as pd
from typing import Dict, Any, List, Tuple


class DocumentProcessor:
    """Base class for document processing."""
    
    def __init__(self, file_path: str):
        """Initialize with file path."""
        self.file_path = file_path
        self.content = None
        
    def extract_text(self) -> str:
        """Extract text from document. To be implemented by subclasses."""
        raise NotImplementedError("Subclasses must implement extract_text()")
    
    def get_metadata(self) -> Dict[str, Any]:
        """Get metadata from document. To be implemented by subclasses."""
        raise NotImplementedError("Subclasses must implement get_metadata()")


class PDFProcessor(DocumentProcessor):
    """Processor for PDF documents."""
    
    def extract_text(self) -> str:
        """Extract text from PDF document."""
        text = ""
        with open(self.file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    
    def get_metadata(self) -> Dict[str, Any]:
        """Get metadata from PDF document."""
        metadata = {}
        with open(self.file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            if pdf_reader.metadata:
                for key, value in pdf_reader.metadata.items():
                    metadata[key] = value
        return metadata
    
    def get_page_count(self) -> int:
        """Get the number of pages in the PDF."""
        with open(self.file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            return len(pdf_reader.pages)


class WordProcessor(DocumentProcessor):
    """Processor for Word documents."""
    
    def extract_text(self) -> str:
        """Extract text from Word document."""
        doc = docx.Document(self.file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    
    def get_metadata(self) -> Dict[str, Any]:
        """Get metadata from Word document."""
        doc = docx.Document(self.file_path)
        metadata = {}
        prop_names = ['author', 'category', 'comments', 'content_status', 
                      'created', 'identifier', 'keywords', 'language', 
                      'last_modified_by', 'last_printed', 'modified', 
                      'revision', 'subject', 'title', 'version']
        
        for prop_name in prop_names:
            if hasattr(doc.core_properties, prop_name):
                prop_value = getattr(doc.core_properties, prop_name)
                if prop_value:
                    metadata[prop_name] = str(prop_value)
        
        return metadata


class ExcelProcessor(DocumentProcessor):
    """Processor for Excel documents."""
    
    def extract_text(self) -> str:
        """Extract text from Excel document."""
        workbook = openpyxl.load_workbook(self.file_path, data_only=True)
        text = ""
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text += f"Sheet: {sheet_name}\n"
            for row in sheet.rows:
                row_text = " | ".join(str(cell.value) if cell.value is not None else "" for cell in row)
                text += row_text + "\n"
            text += "\n"
        return text
    
    def get_metadata(self) -> Dict[str, Any]:
        """Get metadata from Excel document."""
        workbook = openpyxl.load_workbook(self.file_path)
        metadata = {
            'sheet_names': workbook.sheetnames,
            'sheet_count': len(workbook.sheetnames)
        }
        return metadata
    
    def to_dataframe(self) -> Dict[str, pd.DataFrame]:
        """Convert Excel sheets to pandas DataFrames."""
        return pd.read_excel(self.file_path, sheet_name=None)


def get_processor_for_file(file_path: str) -> DocumentProcessor:
    """Factory function to get the appropriate processor for a file."""
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    
    if ext == '.pdf':
        return PDFProcessor(file_path)
    elif ext in ['.docx', '.doc']:
        return WordProcessor(file_path)
    elif ext in ['.xlsx', '.xls']:
        return ExcelProcessor(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def extract_document_sections(text: str) -> List[Tuple[str, str]]:
    """
    Extract sections from a document text.
    Returns a list of tuples (section_title, section_content).
    
    This is a simplified implementation that would need to be adapted
    based on the actual structure of the documents being processed.
    """
    # This is a placeholder implementation
    sections = []
    lines = text.split('\n')
    current_section = None
    current_content = []
    
    for line in lines:
        # Simple heuristic - section titles are uppercase and not too long
        if line.strip().isupper() and len(line.strip()) > 3 and len(line.strip()) < 50:
            if current_section:
                sections.append((current_section, '\n'.join(current_content)))
            current_section = line.strip()
            current_content = []
        else:
            if current_section:
                current_content.append(line)
    
    # Add the last section
    if current_section:
        sections.append((current_section, '\n'.join(current_content)))
    
    return sections 